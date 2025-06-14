import io
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage


def export_resultado(resultado, fig, formato):
    """Return bytes, content type and filename for given format."""
    buffer = io.BytesIO()
    if formato == 'pdf':
        buffer.write(fig.to_image(format='pdf'))
        content_type = 'application/pdf'
        filename = 'resultado.pdf'
    elif formato == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.append(['x1', 'x2', 'Z'])
        ws.append([resultado['x'], resultado['y'], resultado['z']])
        try:
            img_stream = io.BytesIO(fig.to_image(format='png'))
            img = XLImage(img_stream)
            img.anchor = 'E2'
            ws.add_image(img)
        except Exception:
            pass
        wb.save(buffer)
        content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        filename = 'resultado.xlsx'
    elif formato == 'word':
        doc = Document()
        doc.add_heading('Resultado del problema PL', level=1)
        doc.add_paragraph(f"x₁ = {resultado['x']:.2f}")
        doc.add_paragraph(f"x₂ = {resultado['y']:.2f}")
        doc.add_paragraph(f"Z = {resultado['z']:.2f}")
        try:
            img_stream = io.BytesIO(fig.to_image(format='png'))
            doc.add_picture(img_stream, width=Inches(5))
        except Exception:
            pass
        doc.save(buffer)
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = 'resultado.docx'
    else:
        raise ValueError('Formato no soportado')
    buffer.seek(0)
    return buffer.read(), content_type, filename
